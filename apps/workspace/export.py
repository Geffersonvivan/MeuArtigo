"""Exportação do artigo em Markdown, DOCX e PDF.

Resolve os marcadores [[ref:ID]] para citação autor-data + seção Referências (ABNT).
Se o artigo não estiver 'final', prefixa "RASCUNHO — não citar".
"""

from __future__ import annotations

import io
import re

import markdown as md_lib

from apps.memory.citations import gerar_notas_md, gerar_referencias_md, resolver_marcadores

#: perfis de layout do export (fonte / tamanho / entrelinha).
LAYOUT_PROFILES = {
    "abnt":      {"docx_font": "Times New Roman", "pdf_font": "Times, serif",              "size": 12, "lh": 1.5},
    "editorial": {"docx_font": "Georgia",         "pdf_font": "Georgia, serif",            "size": 11, "lh": 1.4},
    "web":       {"docx_font": "Arial",           "pdf_font": "Arial, Helvetica, sans-serif", "size": 14, "lh": 1.6},
}


def _layout(article) -> dict:
    return LAYOUT_PROFILES.get(getattr(article, "perfil_layout", "abnt") or "abnt", LAYOUT_PROFILES["abnt"])


def resolved_markdown(article) -> str:
    refs = {r.pk: r for r in article.references.filter(verificada="ok")}
    modo = getattr(article, "estilo_citacao", "autor_data") or "autor_data"
    numeros: dict[int, int] | None = {} if modo == "nota_rodape" else None
    partes = [f"# {article.titulo}"]
    usados = []
    for sec in article.sections.all():
        partes.append(f"## {sec.titulo}")
        resolvido, u = resolver_marcadores(sec.render_corpo(), refs, modo=modo, numeros=numeros)
        usados.extend(u)
        if resolvido.strip():
            partes.append(resolvido.strip())
    md = "\n\n".join(partes).strip() + "\n"
    if modo == "nota_rodape" and numeros:
        # notas numeradas na ordem de aparição
        ordem = [refs[rid] for rid, _ in sorted(numeros.items(), key=lambda kv: kv[1]) if rid in refs]
        md += "\n\n" + gerar_notas_md(ordem)
    elif usados:
        md += "\n\n" + gerar_referencias_md(usados)
    if article.status != "final":
        md = "> **RASCUNHO — não citar**\n\n" + md
    return md


def to_markdown(article) -> bytes:
    return resolved_markdown(article).encode("utf-8")


def to_docx(article) -> bytes:
    from docx import Document
    from docx.shared import Pt

    lay = _layout(article)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = lay["docx_font"]
    normal.font.size = Pt(lay["size"])
    normal.paragraph_format.line_spacing = lay["lh"]
    for raw in resolved_markdown(article).splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(re.sub(r"[*_>]", "", line[2:]).strip())
            r.italic = True
            r.font.size = Pt(11)
        else:
            doc.add_paragraph(re.sub(r"\*\*(.+?)\*\*", r"\1", line))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(article) -> bytes:
    from xhtml2pdf import pisa

    lay = _layout(article)
    corpo = md_lib.markdown(resolved_markdown(article), extensions=["extra"])
    html = f"""<html><head><meta charset="utf-8"><style>
      @page {{ margin: 2.5cm; }}
      body {{ font-family: {lay['pdf_font']}; font-size: {lay['size']}pt; line-height: {lay['lh']}; color: #111; }}
      h1 {{ font-size: 20pt; }} h2 {{ font-size: 15pt; margin-top: 18pt; }}
      blockquote {{ color: #b45309; font-weight: bold; }}
      p {{ text-align: justify; }}
    </style></head><body>{corpo}</body></html>"""
    buf = io.BytesIO()
    pisa.CreatePDF(src=html, dest=buf, encoding="utf-8")
    return buf.getvalue()


EXPORTERS = {
    "md": (to_markdown, "text/markdown", "md"),
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "pdf": (to_pdf, "application/pdf", "pdf"),
}
